#!/usr/bin/env python3
"""
Test script for enhanced Word document formatting.
"""

import os
import sys
import tempfile

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the idp_common package to the path
sys.path.insert(0, "lib/idp_common_pkg")
from idp_common.ocr.document_converter import DocumentConverter


def create_test_word_document():
    """Create a test Word document with various formatting."""
    doc = WordDocument()

    # Title (Heading 1)
    title = doc.add_heading("Enhanced Formatting Test Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle (Heading 2)
    doc.add_heading("Testing Multi-Level Headings", level=1)

    # Introduction paragraph with mixed formatting
    intro = doc.add_paragraph()
    intro.add_run("This document tests the ").bold = False
    intro.add_run("enhanced formatting capabilities").bold = True
    intro.add_run(" of the DocumentConverter. It includes ")
    intro.add_run("bold text").bold = True
    intro.add_run(", ")
    intro.add_run("italic text").italic = True
    intro.add_run(", and ")
    intro.add_run("underlined text").underline = True
    intro.add_run(".")

    # Level 3 heading
    doc.add_heading("Level 3 Heading", level=3)

    # Centered paragraph
    centered = doc.add_paragraph("This paragraph is center-aligned")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Level 4 heading
    doc.add_heading("Level 4 Heading", level=4)

    # Table with headers
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Product"
    hdr_cells[1].text = "Price"
    hdr_cells[2].text = "Status"

    # Data rows
    data = [
        ("Laptop", "$999", "Available"),
        ("Mouse", "$29", "In Stock"),
        ("Keyboard", "$79", "Limited"),
    ]

    for product, price, status in data:
        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = price
        row_cells[2].text = status

    # Final paragraph
    doc.add_paragraph(
        "This document demonstrates enhanced Word formatting preservation."
    )

    # Save to bytes
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)

        with open(tmp.name, "rb") as f:
            word_bytes = f.read()

        os.unlink(tmp.name)
        return word_bytes


def test_enhanced_formatting():
    """Test the enhanced Word document formatting."""
    print("🔧 Testing Enhanced Word Document Formatting")
    print("=" * 50)

    # Create test document
    print("📄 Creating test Word document...")
    word_bytes = create_test_word_document()
    print(f"   Document size: {len(word_bytes):,} bytes")

    # Initialize DocumentConverter
    print("🔧 Initializing DocumentConverter...")
    converter = DocumentConverter(dpi=150)
    print(f"   Page size: {converter.page_width}x{converter.page_height}")

    # Test font loading
    print("🔤 Testing font loading...")
    try:
        fonts = converter._load_fonts()
        print(f"   ✅ Loaded {len(fonts)} font sizes:")
        for font_name, font_obj in fonts.items():
            print(f"      {font_name}: {font_obj}")
    except Exception as e:
        print(f"   ❌ Font loading failed: {e}")
        return False

    # Process Word document
    print("📝 Processing Word document with enhanced formatting...")
    try:
        pages = converter.convert_word_to_pages(word_bytes)
        print(f"   ✅ Generated {len(pages)} pages")

        for i, (image_bytes, text_content) in enumerate(pages):
            print(
                f"   Page {i + 1}: {len(image_bytes):,} bytes image, {len(text_content)} chars text"
            )

            # Save test image
            test_image_path = f"test_word_page_{i + 1}.jpg"
            with open(test_image_path, "wb") as f:
                f.write(image_bytes)
            print(f"   💾 Saved test image: {test_image_path}")

            # Show text preview
            print(f"   📝 Text preview: {text_content[:100]}...")

        return True

    except Exception as e:
        print(f"   ❌ Processing failed: {e}")
        import traceback

        traceback.print_exc()
        return False


if __name__ == "__main__":
    success = test_enhanced_formatting()
    if success:
        print("\n✅ Enhanced formatting test completed successfully!")
        print(
            "📸 Check the generated test_word_page_*.jpg files to see the enhanced formatting."
        )
    else:
        print("\n❌ Enhanced formatting test failed!")
        sys.exit(1)
